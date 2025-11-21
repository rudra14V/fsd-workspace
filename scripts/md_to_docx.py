import sys
from docx import Document

md_path = sys.argv[1] if len(sys.argv) > 1 else 'chesshive-react/docs/Website_Flow_Telugu.md'
docx_path = sys.argv[2] if len(sys.argv) > 2 else 'chesshive-react/docs/Website_Flow_Telugu.docx'

with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

# Simple conversion: convert headings and paragraphs
doc = Document()
for line in lines:
    s = line.strip('\n')
    if s.startswith('# '):
        doc.add_heading(s[2:].strip(), level=1)
    elif s.startswith('## '):
        doc.add_heading(s[3:].strip(), level=2)
    elif s.startswith('- '):
        # bullet list handling: naive grouping
        doc.add_paragraph(s[2:].strip(), style='List Bullet')
    elif s == '':
        doc.add_paragraph('')
    else:
        doc.add_paragraph(s)

doc.save(docx_path)
print('WROTE', docx_path)
